from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/communication_incentives_proposal.docx")


SECTIONS = [
    (
        "Abstract",
        [
            "This project studies how communication and incentive structure affect the ability of multiple agents to solve a hidden black-box optimization problem. Agents search in a bounded continuous space containing hidden Gaussian peaks. The score of an agent is based only on the hidden value of the point it selects, not on explicit diversity. The central question is whether agents reach better global optimization outcomes under cooperation or competition, and under what communication conditions they approach the global optimum. Diversity is not treated as a reward in the main scoring model. Instead, it is measured as a diagnostic variable that helps explain whether agents are exploring broadly, converging prematurely, or duplicating one another's search effort.",
        ],
    ),
    (
        "1. Motivation",
        [
            "Multi-agent AI systems are increasingly used for search, planning, design, scientific ideation, and other open-ended optimization tasks. A natural assumption is that communication should improve performance: if one agent discovers a high-value region, other agents can learn from that discovery. However, communication can also create premature convergence. If all agents observe the same successful example and imitate it, the system may stop exploring alternative regions of the search space.",
            "This project asks when communication helps or hurts multi-agent optimization. The goal is not simply to measure whether agents are individually intelligent. The goal is to understand how different social conditions, especially cooperation and competition, change the way agents use information and whether the overall system reaches a high-value solution.",
        ],
    ),
    (
        "2. Research Question",
        [
            "The main research question is: Under what communication and incentive conditions do multi-agent systems reach the global optimum in a hidden black-box landscape?",
            "The project compares two incentive settings. In the cooperative setting, agents are instructed to help the system reach the best possible global outcome. In the competitive setting, each agent is instructed to maximize its own score and outperform the others. Both settings use the same hidden landscape, but they differ in how agents interpret communication and whether they treat other agents as collaborators or competitors.",
        ],
    ),
    (
        "3. Environment",
        [
            "The environment is a bounded d-dimensional continuous space. In the default version, d = 10, and each agent submits a vector z_i in [0, 100]^d.",
            "The environment contains K hidden Gaussian peaks. Each peak m has a center mu_m, height h_m, and width sigma_m. The hidden value of a point is V(z_i) = max_m h_m exp(-||z_i - mu_m||_2^2 / (2 sigma_m^2)).",
            "The peak parameters are known to the evaluator but hidden from the agents. Since the evaluator knows all peak heights, the true global optimum is V* = max_m h_m. This makes it possible to measure how close the agents come to the global optimum.",
        ],
    ),
    (
        "4. Scoring Model",
        [
            "The main scoring model is value-only: S_i = V(z_i). This is an intentional change from a diversity-rewarded score. Diversity is not directly rewarded because diversity by itself may not have practical value in a real optimization task. In practice, the goal is usually to find high-value solutions, not to be different for its own sake.",
            "However, diversity remains important as a diagnostic. If many agents collapse to the same region, the system may waste search capacity and miss other high-value peaks. If agents maintain some spread during search, they may discover better solutions. Therefore, diversity is measured to explain search behavior, but it is not part of the reward.",
        ],
    ),
    (
        "5. Incentive Conditions",
        [
            "Cooperative setting: Agents are told that the group should find the best possible solution. The primary objective is system-level optimization. Useful outcomes include discovering the highest peak, reducing the optimality gap, and reaching high value quickly. Communication is expected to be beneficial when it helps agents share discoveries without causing unnecessary duplication.",
            "Competitive setting: Each agent tries to obtain the highest individual score and outperform the others. The environment score is still S_i = V(z_i), but the social objective changes. Agents may prefer to read information from others while sharing less about their own discoveries. This setting tests whether competition encourages useful exploration or whether it reduces information flow and harms global performance.",
        ],
    ),
    (
        "6. Communication Mechanism",
        [
            "Communication is modeled as a two-sided decision: agents choose both what information to reveal and what information to attend to.",
            "In each round, after receiving the score for its previous submission, every agent first decides whether to make its information visible to others. An agent may keep its information private, reveal it to one selected agent, reveal it to a small group, or reveal it to all other agents. The shared information may include the point the agent submitted, the score it received, and a short explanation of its current search strategy.",
            "However, visibility does not mean that the information is automatically used by every recipient. After seeing which information is available, each agent also decides whether to inspect the information shared by others. If many agents make their information visible, a recipient may choose only a subset to read. This attention decision captures the fact that communication has limited value when agents ignore, distrust, or strategically avoid the information provided by others.",
            "The benchmark can also include an explicit negotiation phase between visibility and inspection. In this phase, agents may request information from selected peers, offer reciprocal exchange, and decide whether to accept or reject incoming requests. For example, an agent may ask another agent to reveal its best recent point, offer to exchange its own recent point in return, or refuse to share if it believes the information is strategically valuable. This makes communication an active bargaining process rather than only passive observation.",
            "This two-sided mechanism makes communication strategic. In the cooperative condition, agents may reveal useful information widely and pay attention to discoveries made by others, because the group is trying to find the best possible solution together. In the competitive condition, agents may hide high-value discoveries, selectively reveal information, or choose not to read information from competitors if they suspect it may be misleading or unhelpful.",
            "The evaluator records both sides of communication behavior: how much information agents make visible, how much information other agents actually inspect, whether high-performing agents become more secretive, whether agents tend to attend more to successful agents over time, and whether agents accept or reject information-exchange requests. This makes it possible to distinguish between information availability, information use, and negotiated information exchange.",
        ],
    ),
    (
        "7. Experimental Design",
        [
            "The experiments compare how different incentive and communication conditions affect multi-agent search performance. All experiments use the same hidden black-box landscape, so differences in performance can be attributed to changes in incentive structure, communication structure, or landscape difficulty.",
            "Experiment 1: Cooperation versus Competition. The first experiment compares cooperative and competitive agents under the same communication mechanism. In the cooperative condition, agents are instructed to help the group find the best possible solution. In the competitive condition, each agent is instructed to maximize its own score and outperform the other agents. This experiment tests whether cooperative incentives improve system-level optimization by encouraging information sharing, or whether competitive incentives sometimes produce better exploration because agents are less likely to imitate one another.",
            "Experiment 2: Communication Structure. The second experiment varies how much information can move between agents. The main conditions are no communication, full visibility, and strategic visibility. In the no-communication condition, agents only observe their own previous submissions and scores. In the full-visibility condition, all agents can see all previous submissions and scores from every other agent. In the strategic-visibility condition, each agent decides which other agents can see its information, and each recipient decides whether to inspect the information made available to it.",
            "This experiment tests whether full communication helps agents discover high-value regions faster, or whether it causes premature convergence. It also tests whether strategic communication can balance the benefits of information sharing with the need for independent exploration.",
            "Experiment 3: Negotiated Information Exchange. The third experiment adds an explicit negotiation phase. Agents may request information from selected peers, offer reciprocal exchange, accept requests, or reject requests. This condition is compared against passive strategic visibility, where agents only decide what to reveal and what to inspect. This experiment tests whether negotiation improves global optimization by allowing agents to target useful information, or whether it creates strategic withholding in competitive settings. It also tests whether cooperative agents use negotiation to coordinate search more effectively than competitive agents.",
            "Experiment 4: Landscape Difficulty. The fourth experiment varies the difficulty of the hidden landscape. Difficulty is changed by adjusting the number of peaks, the width of the peaks, and the difference between peak heights. Broad peaks create easier search problems because agents are more likely to find high-value regions by chance. Narrow peaks create harder search problems because agents must search more precisely. Landscapes with many similar peaks test whether agents spread across multiple promising regions, while landscapes with one dominant peak test whether agents can efficiently converge on the global optimum.",
            "Experiment 5: Population Size. The fifth experiment varies the number of agents. A larger population may improve coverage of the search space because more points are explored in each round. However, a larger population may also increase imitation if agents communicate too much and repeatedly follow the same visible high-scoring examples. This experiment tests when additional agents improve global optimization and when they mainly create redundant search.",
        ],
    ),
    (
        "8. Evaluation Metrics",
        [
            "The primary metrics are value-based: best value found by any agent across all rounds; optimality gap, V* - max_{i,t} V(z_{i,t}); time to reach a fixed percentage of the global optimum; average final value across agents; and success rate across random seeds.",
            "Communication metrics distinguish between information availability, information use, and negotiated exchange: mean number of agents that each agent makes its information visible to; frequency of private decisions, where an agent chooses to reveal information to no one; mean number of visible peer observations that each agent actually inspects; mean observed peer count; whether high-performing agents become more or less willing to share over time; request count; acceptance rate; rejection rate; reciprocal exchange rate; information asymmetry; and communication efficiency, measured as performance improvement per unit of inspected information.",
            "Diversity and convergence are diagnostic metrics: average pairwise distance between agents' submitted points, peak coverage, agent clustering near the same peak, and convergence rate over rounds. These diagnostic metrics are not rewards. They are used to explain why agents succeed or fail.",
        ],
    ),
    (
        "9. Practical Implications",
        [
            "The practical implication is that communication policy matters in multi-agent AI systems. If many agents are assigned to solve the same problem, unrestricted information sharing may cause them to duplicate one another. On the other hand, no communication may waste useful discoveries. The best system may require structured or strategic communication: agents should share enough information to improve search, but not so much that the whole population collapses into the same behavior.",
            "This applies to AI systems used for scientific research, design exploration, planning, and automated problem solving. In these settings, the practical goal is not diversity itself. The goal is high-quality optimization. Diversity matters only when it helps the system avoid redundant search and discover better solutions.",
        ],
    ),
    (
        "10. Contribution",
        [
            "This benchmark contributes a controlled framework for studying communication and incentives in multi-agent black-box optimization. It separates three concepts that are often mixed together: objective value, social incentive, and diversity. The objective value is the hidden score that agents try to optimize. The social incentive determines whether agents behave cooperatively or competitively. Diversity is measured only as a diagnostic factor that explains search dynamics.",
            "The benchmark therefore asks a concrete question: When do multiple communicating agents reach the global optimum, and how does the answer change under cooperation versus competition?",
        ],
    ),
    (
        "11. Conclusion",
        [
            "This project studies multi-agent optimization in a hidden peak landscape. Agents search for high-value solutions while choosing visibility, requests, reciprocal exchange, and accept/reject behavior. The main score is value-only, making the benchmark closer to practical optimization tasks. The project compares cooperative and competitive incentives and evaluates which communication conditions help agents approach the global optimum. Diversity is not treated as a reward, but it remains useful as a diagnostic for understanding exploration, convergence, and redundant search.",
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(46, 116, 181)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(46, 116, 181)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Communication and Incentives in Multi-Agent")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)
    title.add_run().add_break()
    run = title.add_run("Black-Box Optimization")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Revised Research Proposal")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_overview_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Design Choice"
    hdr[1].text = "Current Proposal"
    for cell in hdr:
        set_cell_shading(cell, "F4F6F9")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    rows = [
        ("Main objective", "Reach the highest hidden value / global optimum."),
        ("Agent incentives", "Compare cooperative and competitive settings."),
        ("Reward", "Value-only: S_i = V(z_i)."),
        ("Diversity", "Diagnostic only; not part of the reward."),
        ("Communication", "Agents choose visibility and inspection decisions each round."),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    apply_document_styles(doc)
    add_title(doc)
    add_overview_table(doc)

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
